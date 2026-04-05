"""
Student CRUD API — Phase 2 complete.

ISOLATION GUARANTEE:
- agency_id ALWAYS comes from JWT via get_current_user().
- It is NEVER accepted from request body, query params, or URL.
- Every query filters by user.agency_id before touching the DB.
"""

import base64
import json as _json

from fastapi import APIRouter, Depends, HTTPException, Request, Query, UploadFile, File
from pydantic import BaseModel, Field, field_validator, EmailStr
from typing import Optional, List, Any
from datetime import date, datetime, timezone

from core.auth import get_current_user
from core.db import get_service_client
from core.audit import write_audit_log
from models.user import AuthUser

router = APIRouter(tags=["Students"])


# ──────────────────────────────────────────────────────────────
# Pydantic models — mirror the full DB schema exactly
# ──────────────────────────────────────────────────────────────

class StudentCreate(BaseModel):
    # Core identity
    full_name: str = Field(min_length=1, max_length=200)
    preferred_name: Optional[str] = Field(default=None, max_length=100)
    date_of_birth: Optional[date] = None
    nationality: Optional[str] = Field(default=None, max_length=100)
    email: Optional[EmailStr] = None
    phone: Optional[str] = Field(default=None, max_length=30)
    telegram_username: Optional[str] = Field(default=None, max_length=64)
    status: str = Field(default="intake", max_length=50)
    season: Optional[str] = Field(default=None, max_length=20)
    assigned_staff_id: Optional[str] = None
    photo_url: Optional[str] = Field(default=None, max_length=2048)

    @field_validator("full_name", mode="before")
    @classmethod
    def strip_full_name(cls, v: Any) -> Any:
        if isinstance(v, str):
            v = v.strip()
            if not v:
                raise ValueError("full_name cannot be blank")
        return v

    # Personal background (Common App: Personal Info)
    gender: Optional[str] = None
    pronouns: Optional[str] = None
    city_of_birth: Optional[str] = None
    country_of_birth: Optional[str] = None
    visa_status: Optional[str] = None
    languages_at_home: Optional[str] = None

    # Passport & travel
    passport_number: Optional[str] = None
    passport_expiry: Optional[date] = None
    languages: List[Any] = []

    # Address
    address_street: Optional[str] = None
    address_city: Optional[str] = None
    address_country: Optional[str] = None
    address_zip: Optional[str] = None

    # Parent — kept for backwards compat; prefer father_*/mother_* below
    parent_name: Optional[str] = None
    parent_email: Optional[str] = None
    parent_phone: Optional[str] = None

    # Father (Common App: Family)
    father_name: Optional[str] = None
    father_email: Optional[str] = None
    father_phone: Optional[str] = None
    father_education: Optional[str] = None
    father_occupation: Optional[str] = None
    father_employer: Optional[str] = None

    # Mother (Common App: Family)
    mother_name: Optional[str] = None
    mother_email: Optional[str] = None
    mother_phone: Optional[str] = None
    mother_education: Optional[str] = None
    mother_occupation: Optional[str] = None
    mother_employer: Optional[str] = None

    # Family context
    parents_marital_status: Optional[str] = None
    first_generation_student: Optional[bool] = False

    # Academic record (Common App: Education)
    high_school_name: Optional[str] = Field(default=None, max_length=200)
    high_school_country: Optional[str] = Field(default=None, max_length=100)
    school_ceeb_code: Optional[str] = Field(default=None, max_length=20)
    school_type: Optional[str] = Field(default=None, max_length=50)
    school_city: Optional[str] = Field(default=None, max_length=100)
    graduation_year: Optional[int] = Field(default=None, ge=2000, le=2040)
    gpa: Optional[float] = Field(default=None, ge=0.0, le=5.0)
    gpa_scale: float = Field(default=4.0, ge=1.0, le=5.0)
    class_rank: Optional[str] = Field(default=None, max_length=50)
    class_size: Optional[int] = Field(default=None, ge=1, le=10000)

    # Standardized test scores
    sat_total: Optional[int] = Field(default=None, ge=400, le=1600)
    sat_math: Optional[int] = Field(default=None, ge=200, le=800)
    sat_reading: Optional[int] = Field(default=None, ge=200, le=800)
    sat_essay: Optional[int] = None
    act_score: Optional[int] = Field(default=None, ge=1, le=36)
    act_english: Optional[int] = Field(default=None, ge=1, le=36)
    act_math: Optional[int] = Field(default=None, ge=1, le=36)
    act_reading: Optional[int] = Field(default=None, ge=1, le=36)
    act_science: Optional[int] = Field(default=None, ge=1, le=36)
    toefl_score: Optional[int] = Field(default=None, ge=0, le=120)
    ielts_score: Optional[float] = Field(default=None, ge=0.0, le=9.0)
    duolingo_score: Optional[int] = Field(default=None, ge=10, le=160)
    ap_scores: List[Any] = []
    ib_scores: List[Any] = []

    # Activities, honors, work
    activities: List[Any] = []
    awards: List[Any] = []
    work_experience: List[Any] = []

    # Application intent
    intended_major: Optional[str] = Field(default=None, max_length=200)
    application_type: str = Field(default="freshman", max_length=50)

    # Recommendation context
    teacher_rec_info: List[Any] = []
    counselor_notes: Optional[str] = Field(default=None, max_length=10000)

    # Internal
    notes: Optional[str] = Field(default=None, max_length=10000)


class StudentUpdate(BaseModel):
    # Core identity
    full_name: Optional[str] = Field(default=None, min_length=1, max_length=200)
    preferred_name: Optional[str] = Field(default=None, max_length=100)
    date_of_birth: Optional[date] = None
    nationality: Optional[str] = Field(default=None, max_length=100)
    email: Optional[EmailStr] = None
    phone: Optional[str] = Field(default=None, max_length=30)
    telegram_username: Optional[str] = Field(default=None, max_length=64)
    status: Optional[str] = Field(default=None, max_length=50)
    season: Optional[str] = Field(default=None, max_length=20)
    assigned_staff_id: Optional[str] = None
    photo_url: Optional[str] = Field(default=None, max_length=2048)

    # Personal background
    gender: Optional[str] = Field(default=None, max_length=50)
    pronouns: Optional[str] = Field(default=None, max_length=50)
    city_of_birth: Optional[str] = Field(default=None, max_length=100)
    country_of_birth: Optional[str] = Field(default=None, max_length=100)
    visa_status: Optional[str] = Field(default=None, max_length=50)
    languages_at_home: Optional[str] = Field(default=None, max_length=200)

    # Passport & travel
    passport_number: Optional[str] = Field(default=None, max_length=20)
    passport_expiry: Optional[date] = None
    languages: Optional[List[Any]] = None

    # Address
    address_street: Optional[str] = Field(default=None, max_length=200)
    address_city: Optional[str] = Field(default=None, max_length=100)
    address_country: Optional[str] = Field(default=None, max_length=100)
    address_zip: Optional[str] = Field(default=None, max_length=20)

    # Parent (backwards compat)
    parent_name: Optional[str] = Field(default=None, max_length=200)
    parent_email: Optional[EmailStr] = None
    parent_phone: Optional[str] = Field(default=None, max_length=30)

    # Father
    father_name: Optional[str] = Field(default=None, max_length=200)
    father_email: Optional[EmailStr] = None
    father_phone: Optional[str] = Field(default=None, max_length=30)
    father_education: Optional[str] = Field(default=None, max_length=100)
    father_occupation: Optional[str] = Field(default=None, max_length=100)
    father_employer: Optional[str] = Field(default=None, max_length=200)

    # Mother
    mother_name: Optional[str] = Field(default=None, max_length=200)
    mother_email: Optional[EmailStr] = None
    mother_phone: Optional[str] = Field(default=None, max_length=30)
    mother_education: Optional[str] = Field(default=None, max_length=100)
    mother_occupation: Optional[str] = Field(default=None, max_length=100)
    mother_employer: Optional[str] = Field(default=None, max_length=200)

    # Family context
    parents_marital_status: Optional[str] = Field(default=None, max_length=50)
    first_generation_student: Optional[bool] = None

    # Academic record
    high_school_name: Optional[str] = Field(default=None, max_length=200)
    high_school_country: Optional[str] = Field(default=None, max_length=100)
    school_ceeb_code: Optional[str] = Field(default=None, max_length=20)
    school_type: Optional[str] = Field(default=None, max_length=50)
    school_city: Optional[str] = Field(default=None, max_length=100)
    graduation_year: Optional[int] = Field(default=None, ge=2000, le=2040)
    gpa: Optional[float] = Field(default=None, ge=0.0, le=5.0)
    gpa_scale: Optional[float] = Field(default=None, ge=1.0, le=5.0)
    class_rank: Optional[str] = Field(default=None, max_length=50)
    class_size: Optional[int] = Field(default=None, ge=1, le=10000)

    # Test scores
    sat_total: Optional[int] = Field(default=None, ge=400, le=1600)
    sat_math: Optional[int] = Field(default=None, ge=200, le=800)
    sat_reading: Optional[int] = Field(default=None, ge=200, le=800)
    sat_essay: Optional[int] = None
    act_score: Optional[int] = Field(default=None, ge=1, le=36)
    act_english: Optional[int] = Field(default=None, ge=1, le=36)
    act_math: Optional[int] = Field(default=None, ge=1, le=36)
    act_reading: Optional[int] = Field(default=None, ge=1, le=36)
    act_science: Optional[int] = Field(default=None, ge=1, le=36)
    toefl_score: Optional[int] = Field(default=None, ge=0, le=120)
    ielts_score: Optional[float] = Field(default=None, ge=0.0, le=9.0)
    duolingo_score: Optional[int] = Field(default=None, ge=10, le=160)
    ap_scores: Optional[List[Any]] = None
    ib_scores: Optional[List[Any]] = None

    # Unified English proficiency (replaces toefl/ielts/duolingo separate fields)
    english_test_type: Optional[str] = None    # 'toefl_ibt' | 'ielts' | 'duolingo' | 'pte' | 'cambridge'
    english_test_scores: Optional[dict] = None  # all sub-scores as JSON

    # Activities, honors, work
    activities: Optional[List[Any]] = None
    awards: Optional[List[Any]] = None
    work_experience: Optional[List[Any]] = None

    # Application intent
    intended_major: Optional[str] = Field(default=None, max_length=200)
    application_type: Optional[str] = Field(default=None, max_length=50)

    # Recommendation context
    teacher_rec_info: Optional[List[Any]] = None
    counselor_notes: Optional[str] = Field(default=None, max_length=10000)

    # Internal
    notes: Optional[str] = Field(default=None, max_length=10000)


class BulkImportRequest(BaseModel):
    students: List[dict]


# ──────────────────────────────────────────────────────────────
# Routes
# ──────────────────────────────────────────────────────────────

@router.get("/students")
async def list_students(
    q: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    season: Optional[str] = Query(None),
    include_archived: bool = Query(False),
    limit: int = Query(50, le=200),
    offset: int = Query(0, ge=0),
    user: AuthUser = Depends(get_current_user),
):
    db = get_service_client()

    query = (
        db.table("students")
        .select(
            "id, full_name, preferred_name, email, status, season, "
            "graduation_year, gpa, gpa_scale, sat_total, act_score, "
            "intended_major, nationality, high_school_name, "
            "assigned_staff_id, created_at, updated_at, "
            "applications(id, university_name, status, deadline_regular)",
            count="exact",
        )
        .eq("agency_id", user.agency_id)  # ISOLATION — always from JWT
    )

    if q:
        query = query.ilike("full_name", f"%{q}%")
    if status:
        # Explicit status filter — show exactly those students
        query = query.eq("status", status)
    elif not include_archived:
        # Default: hide archived students so they don't consume visible slots
        query = query.neq("status", "archived")
    if season:
        query = query.eq("season", season)

    result = (
        query.order("updated_at", desc=True)
        .range(offset, offset + limit - 1)
        .execute()
    )

    return {"students": result.data or [], "total": result.count or 0}


@router.post("/students", status_code=201)
async def create_student(
    data: StudentCreate,
    request: Request,
    user: AuthUser = Depends(get_current_user),
):
    db = get_service_client()

    payload = data.model_dump()
    payload["agency_id"] = user.agency_id  # ISOLATION — overwrite unconditionally

    for f in ("date_of_birth", "passport_expiry"):
        if payload.get(f) and hasattr(payload[f], "isoformat"):
            payload[f] = payload[f].isoformat()

    result = db.table("students").insert(payload).execute()
    if not result.data:
        raise HTTPException(500, "Failed to create student")

    student = result.data[0]

    await write_audit_log(
        agency_id=user.agency_id,
        user_id=user.id,
        student_id=student["id"],
        action="student.created",
        entity_type="student",
        entity_id=student["id"],
        old_value=None,
        new_value={"full_name": student["full_name"], "status": student["status"]},
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )

    return student


@router.get("/students/{student_id}")
async def get_student(
    student_id: str,
    user: AuthUser = Depends(get_current_user),
):
    db = get_service_client()

    result = (
        db.table("students")
        .select("*")
        .eq("id", student_id)
        .eq("agency_id", user.agency_id)  # ISOLATION
        .single()
        .execute()
    )

    if not result.data:
        raise HTTPException(404, "Student not found")

    return result.data


@router.patch("/students/{student_id}")
async def update_student(
    student_id: str,
    data: StudentUpdate,
    request: Request,
    user: AuthUser = Depends(get_current_user),
):
    db = get_service_client()

    existing = (
        db.table("students")
        .select("*")
        .eq("id", student_id)
        .eq("agency_id", user.agency_id)  # ISOLATION
        .single()
        .execute()
    )
    if not existing.data:
        raise HTTPException(404, "Student not found")

    payload = {k: v for k, v in data.model_dump().items() if v is not None}
    payload["updated_at"] = datetime.now(timezone.utc).isoformat()

    for f in ("date_of_birth", "passport_expiry"):
        if payload.get(f) and hasattr(payload[f], "isoformat"):
            payload[f] = payload[f].isoformat()

    # Archive / restore logic
    current_status = existing.data.get("status", "intake")
    if payload.get("status") == "archived" and current_status != "archived":
        # Store the current status so we can restore to it later
        payload["pre_archive_status"] = current_status
    elif payload.get("status") and payload["status"] != "archived" and current_status == "archived":
        # Restoring — use pre_archive_status if available, else keep what was requested
        pre = existing.data.get("pre_archive_status")
        if pre and pre != "archived":
            payload["status"] = pre
        # Clear the stored pre-archive status
        payload["pre_archive_status"] = None

    result = (
        db.table("students")
        .update(payload)
        .eq("id", student_id)
        .eq("agency_id", user.agency_id)  # ISOLATION
        .execute()
    )
    if not result.data:
        raise HTTPException(500, "Update failed")

    changed_old = {k: existing.data.get(k) for k in payload if k != "updated_at"}
    changed_new = {k: payload[k] for k in payload if k != "updated_at"}

    await write_audit_log(
        agency_id=user.agency_id,
        user_id=user.id,
        student_id=student_id,
        action="student.updated",
        entity_type="student",
        entity_id=student_id,
        old_value=changed_old,
        new_value=changed_new,
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )

    return result.data[0]


@router.delete("/students/{student_id}", status_code=204)
async def delete_student(
    student_id: str,
    request: Request,
    user: AuthUser = Depends(get_current_user),
):
    db = get_service_client()

    existing = (
        db.table("students")
        .select("id, full_name, status")
        .eq("id", student_id)
        .eq("agency_id", user.agency_id)  # ISOLATION
        .single()
        .execute()
    )
    if not existing.data:
        raise HTTPException(404, "Student not found")

    student_name = existing.data["full_name"]
    student_status = existing.data["status"]

    # ── Explicit cascade delete (works regardless of DB-level FK settings) ──
    # Order matters: delete leaves before trunks.

    # 1. Null-out audit_logs — keep history but break the FK
    db.table("audit_logs").update({"student_id": None}).eq("student_id", student_id).execute()

    # 2. Essay versions are children of essays — delete essays first (versions cascade)
    essay_ids_res = db.table("essays").select("id").eq("student_id", student_id).execute()
    essay_ids = [r["id"] for r in (essay_ids_res.data or [])]
    if essay_ids:
        db.table("essay_versions").delete().in_("essay_id", essay_ids).execute()
    db.table("essays").delete().eq("student_id", student_id).execute()

    # 3. Delete tables that reference applications(id) without CASCADE,
    #    so deleting applications won't be blocked.
    app_ids_res = db.table("applications").select("id").eq("student_id", student_id).execute()
    app_ids = [r["id"] for r in (app_ids_res.data or [])]
    if app_ids:
        db.table("deadlines").delete().in_("application_id", app_ids).execute()
        db.table("documents").delete().in_("application_id", app_ids).execute()
        db.table("recommendation_letters").delete().in_("application_id", app_ids).execute()

    # 4. Delete remaining child tables (student_id FK)
    for tbl in [
        "applications",
        "deadlines",
        "documents",
        "recommendation_letters",
        "recommenders",
        "agent_jobs",
        "emails",
        "student_credentials",
        "automation_workflows",
    ]:
        try:
            db.table(tbl).delete().eq("student_id", student_id).execute()
        except Exception:
            pass  # Table may not exist or already empty — safe to continue

    # 5. Finally delete the student row
    result = db.table("students").delete().eq("id", student_id).eq(
        "agency_id", user.agency_id
    ).execute()

    if result.data is not None and len(result.data) == 0:
        # Supabase returns [] when nothing was deleted — likely already gone, treat as success
        pass

    await write_audit_log(
        agency_id=user.agency_id,
        user_id=user.id,
        student_id=None,
        action="student.deleted",
        entity_type="student",
        entity_id=student_id,
        old_value={"full_name": student_name, "status": student_status},
        new_value=None,
        ip_address=request.client.host if request.client else None,
        user_agent=request.headers.get("user-agent"),
    )


@router.post("/students/import", status_code=201)
async def bulk_import_students(
    data: BulkImportRequest,
    request: Request,
    user: AuthUser = Depends(get_current_user),
):
    """
    Bulk import students from a JSON array.
    Validates each row has at least full_name.
    Returns created count and any errors.
    """
    db = get_service_client()

    if not data.students:
        return {"created": 0, "errors": []}

    if len(data.students) > 100:
        raise HTTPException(400, "Maximum 100 students per import")

    created_count = 0
    errors = []

    for row_idx, student_data in enumerate(data.students):
        try:
            # Validate required field
            if not isinstance(student_data, dict) or not student_data.get("full_name"):
                errors.append({
                    "row": row_idx + 1,
                    "name": student_data.get("full_name", "unknown") if isinstance(student_data, dict) else "unknown",
                    "error": "full_name is required"
                })
                continue

            # Use StudentCreate to validate the schema
            try:
                validated = StudentCreate(**student_data)
            except Exception as ve:
                errors.append({
                    "row": row_idx + 1,
                    "name": student_data.get("full_name", "unknown"),
                    "error": f"Validation error: {str(ve)[:100]}"
                })
                continue

            # Prepare payload
            payload = validated.model_dump()
            payload["agency_id"] = user.agency_id  # ISOLATION — overwrite unconditionally

            # Convert date fields to ISO format
            for f in ("date_of_birth", "passport_expiry"):
                if payload.get(f) and hasattr(payload[f], "isoformat"):
                    payload[f] = payload[f].isoformat()

            # Insert student
            result = db.table("students").insert(payload).execute()
            if not result.data:
                errors.append({
                    "row": row_idx + 1,
                    "name": student_data.get("full_name", "unknown"),
                    "error": "Database insert failed"
                })
                continue

            student = result.data[0]
            created_count += 1

            # Audit log for this imported student
            await write_audit_log(
                agency_id=user.agency_id,
                user_id=user.id,
                student_id=student["id"],
                action="student.imported",
                entity_type="student",
                entity_id=student["id"],
                old_value=None,
                new_value={"full_name": student["full_name"], "status": student["status"]},
                ip_address=request.client.host if request.client else None,
                user_agent=request.headers.get("user-agent"),
            )

        except Exception as e:
            errors.append({
                "row": row_idx + 1,
                "name": student_data.get("full_name", "unknown") if isinstance(student_data, dict) else "unknown",
                "error": str(e)[:100]
            })

    return {"created": created_count, "errors": errors}


# ──────────────────────────────────────────────────────────────
# English test score extraction via GPT-4o Vision
# ──────────────────────────────────────────────────────────────

_EXTRACTION_PROMPT = """
You are an expert at reading official English language test score reports.
Look at this score report image and extract ALL scores you can see.

Return ONLY valid JSON in exactly this format (no markdown, no explanation):
{
  "test_type": "<one of: toefl_ibt, ielts, duolingo, pte, cambridge>",
  "scores": {
    "total": <number or null>,
    "reading": <number or null>,
    "listening": <number or null>,
    "speaking": <number or null>,
    "writing": <number or null>,
    "literacy": <number or null>,
    "comprehension": <number or null>,
    "conversation": <number or null>,
    "production": <number or null>,
    "use_of_english": <number or null>,
    "grade": "<A/B/C or null>"
  }
}

Rules:
- For TOEFL iBT: total is 0-120, each section is 0-30.
- For IELTS: overall band and each section are 0-9 (half bands like 7.5 are valid).
- For Duolingo: overall and sub-scores are 10-160.
- For PTE Academic: overall and sections are 10-90.
- For Cambridge (C1/C2): scores are 0-100 per component plus a grade (A/B/C).
- If you cannot determine the test type, use your best guess.
- Only include fields that are actually present in the document; set others to null.
""".strip()


@router.post("/students/{student_id}/extract-english-scores")
async def extract_english_scores(
    student_id: str,
    file: UploadFile = File(...),
    user: AuthUser = Depends(get_current_user),
):
    """
    Upload a score report image or PDF.
    GPT-4o Vision extracts all sub-scores and returns structured JSON.
    The caller is responsible for saving the result via PATCH /students/{id}.
    """
    db = get_service_client()

    # Verify student belongs to this agency
    exists = (
        db.table("students")
        .select("id")
        .eq("id", student_id)
        .eq("agency_id", user.agency_id)
        .maybe_single()
        .execute()
    )
    if not exists.data:
        raise HTTPException(404, "Student not found")

    ALLOWED = {"image/jpeg", "image/png", "image/webp", "image/gif", "application/pdf"}
    MAX_BYTES = 10 * 1024 * 1024  # 10 MB

    content_type = file.content_type or "application/octet-stream"
    if content_type not in ALLOWED:
        raise HTTPException(400, f"Unsupported file type: {content_type}. Upload JPEG, PNG, WebP, or PDF.")

    raw = await file.read()
    if len(raw) > MAX_BYTES:
        raise HTTPException(400, "File too large. Maximum size is 10 MB.")

    # For PDFs we send as base64 data URL; GPT-4o handles single-page PDFs fine.
    b64 = base64.b64encode(raw).decode()
    data_url = f"data:{content_type};base64,{b64}"

    from openai import AsyncOpenAI
    from core.config import settings
    client = AsyncOpenAI(api_key=settings.OPENAI_API_KEY)

    try:
        response = await client.chat.completions.create(
            model=settings.AI_MODEL_SMART,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _EXTRACTION_PROMPT},
                        {"type": "image_url", "image_url": {"url": data_url, "detail": "high"}},
                    ],
                }
            ],
            max_tokens=512,
            temperature=0,
        )
    except Exception as e:
        raise HTTPException(502, f"AI extraction failed: {e}")

    raw_text = response.choices[0].message.content.strip()

    # Strip markdown code fences if present
    if raw_text.startswith("```"):
        raw_text = raw_text.split("```")[1]
        if raw_text.startswith("json"):
            raw_text = raw_text[4:]

    try:
        result = _json.loads(raw_text)
    except Exception:
        raise HTTPException(502, f"AI returned unparseable response: {raw_text[:200]}")

    return result


# ── Smart Form Upload ──────────────────────────────────────────────────────────

@router.post("/students/parse-form")
async def parse_intake_form(
    file: UploadFile = File(...),
    user: AuthUser = Depends(get_current_user),
):
    """
    Parse a completed ApplyPilot client intake form (.docx).
    Does NOT create a student record — returns extracted fields for review.
    Use POST /students after the coordinator has reviewed/confirmed.
    """
    if not (file.filename or "").lower().endswith(".docx"):
        raise HTTPException(400, "Only .docx files are accepted")

    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(400, "File too large — maximum 10 MB")

    try:
        from docx import Document
        import io as _io
        doc = Document(_io.BytesIO(content))
    except Exception as exc:
        raise HTTPException(400, f"Could not read .docx file: {exc}")

    # ── Extract all label→value pairs from 2-column tables ────────────────────
    raw: dict[str, str] = {}
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) < 2:
                continue
            # Skip merged/colspan rows (section dividers)
            if cells[0]._tc is cells[1]._tc:
                continue
            label_paras = cells[0].paragraphs
            if not label_paras:
                continue
            label = label_paras[0].text.strip().lstrip("\u25B6").strip().rstrip("*").strip().rstrip(":").strip()
            value = cells[1].text.strip()
            if label and value and value.lower() not in ("n/a", "na", "-", "\u2014", ""):
                raw[label] = value

    # ── Helpers ───────────────────────────────────────────────────────────────
    def _int(v):
        try:
            return int(str(v).replace(",", "").strip())
        except Exception:
            return None

    def _float(v):
        try:
            return float(str(v).strip())
        except Exception:
            return None

    def _year(v):
        try:
            parts = str(v).strip().split("/")
            return int(parts[-1]) if len(parts) >= 2 else int(str(v).strip()[:4])
        except Exception:
            return None

    def _date(v):
        if not v:
            return None
        try:
            p = str(v).strip().split("/")
            if len(p) == 3:
                return f"{p[2]}-{p[1].zfill(2)}-{p[0].zfill(2)}"
        except Exception:
            pass
        return None

    def _get(key):
        v = raw.get(key, "").strip()
        return v or None

    # ── Build mapped output ───────────────────────────────────────────────────
    first = (_get("First Name") or "").strip()
    last  = (_get("Last Name") or "").strip()
    full_name = f"{first} {last}".strip() or None

    father_first = (_get("Father First Name") or "").strip()
    father_last  = (_get("Father Last Name") or "").strip()
    father_name  = f"{father_first} {father_last}".strip() or None

    mother_first = (_get("Mother First Name") or "").strip()
    mother_last  = (_get("Mother Last Name") or "").strip()
    mother_name  = f"{mother_first} {mother_last}".strip() or None

    extracted = {
        "full_name":          full_name,
        "preferred_name":     _get("Preferred / Nick Name"),
        "date_of_birth":      _date(_get("Date of Birth") or ""),
        "gender":             _get("Gender"),
        "nationality":        _get("Country of Citizenship"),
        "country_of_birth":   _get("Country of Birth"),
        "city_of_birth":      _get("City of Birth"),
        "email":              _get("Email Address"),
        "phone":              _get("Phone Number"),
        "address_street":     _get("Street Address"),
        "address_city":       _get("City"),
        "address_country":    _get("Country"),
        "address_zip":        _get("Zip / Postal Code"),
        "passport_number":    _get("Passport Number"),
        "passport_expiry":    _date(_get("Date of Expiration") or ""),
        "father_name":        father_name,
        "father_email":       _get("Father Email Address"),
        "father_phone":       _get("Father Phone Number"),
        "father_education":   _get("Father Education Level"),
        "father_occupation":  _get("Father Occupation"),
        "mother_name":        mother_name,
        "mother_email":       _get("Mother Email Address"),
        "mother_phone":       _get("Mother Phone Number"),
        "mother_education":   _get("Mother Education Level"),
        "mother_occupation":  _get("Mother Occupation"),
        "parent_name":        _get("Emergency Contact Full Name"),
        "parent_phone":       _get("Emergency Contact Phone"),
        "parent_email":       _get("Emergency Contact Email"),
        "high_school_name":   _get("High School Name"),
        "high_school_country": _get("School Country"),
        "school_city":        _get("School City"),
        "graduation_year":    _year(_get("Attendance To") or ""),
        "gpa":                _float(_get("Cumulative GPA") or ""),
        "gpa_scale":          _float(_get("GPA Scale") or "") or 4.0,
        "class_rank":         _get("Class Rank"),
        "sat_total":          _int(_get("SAT Total Score") or ""),
        "sat_math":           _int(_get("SAT Math Score") or ""),
        "sat_reading":        _int(_get("SAT Reading and Writing Score") or ""),
        "act_score":          _int(_get("ACT Composite Score") or ""),
        "toefl_score":        _int(_get("TOEFL Total Score") or ""),
        "ielts_score":        _float(_get("IELTS Band Score") or ""),
        "season":             _get("Intended Enrollment Term"),
        "languages_at_home":  _get("Languages Spoken"),
        "status":             "intake",
    }

    # Remove None values before returning
    extracted = {k: v for k, v in extracted.items() if v is not None}

    return {
        "extracted":       extracted,
        "raw_fields_found": len(raw),
        "mapped_fields":   len(extracted),
        "source_file":     file.filename,
    }
